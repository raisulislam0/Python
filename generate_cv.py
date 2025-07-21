import requests
from docx import Document
from docx.shared import Pt


OLLAMA_URL = "http://localhost:11434/api/generate"

def generate_cv_with_llama(prompt):
    """Generate text from local LLaMA 3.2 via Ollama API"""
    response = requests.post(OLLAMA_URL, json={
        "model": "llama3.2",
        "prompt": prompt,
        "stream": False
    })

    if response.status_code != 200:
        raise Exception(f"Error: {response.status_code} - {response.text}")

    return response.json()["response"]

def write_cv_to_word(content, filename="Generated_CV.docx"):
    """Format CV content with styling and save as a Word document"""
    doc = Document()

    lines = [line.strip() for line in content.strip().split("\n") if line.strip()]
    current_section = None

    # Format first line as name/title
    doc.add_heading(lines[0], level=0)

    for line in lines[1:]:
        # Section headings
        if line.endswith(":"):
            current_section = line.rstrip(":").strip()
            doc.add_heading(current_section, level=1)
        # Bullet list for skills/projects
        elif current_section in ["Skills", "Projects"] and line.startswith("-"):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")
        # Bold + text for Work/Education bullets
        elif current_section in ["Work Experience", "Education"] and line.startswith("-"):
            p = doc.add_paragraph()
            p.add_run(line[1:].strip()).bold = True
        else:
            doc.add_paragraph(line)

    # Set font for all paragraphs
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.save(filename)
    print(f"✅ CV saved as '{filename}'")

def main():
    prompt = """
Write a professional one-page CV for a software engineer named Raisul Islam.
Include the following sections:
- Full Name and Contact Information
- Professional Summary
- Skills (Python, Docker, Git, Redis, LLMs)
- Work Experience at Fiftytwo Digital Ltd
- Education
- Projects (RAG System, AI IDE, Streamlit Chatbot)
Format the CV cleanly with headings, bullet points, and short concise sentences.
Do NOT use markdown.
Just write in structured plain text suitable for formatting into a Word document.
Start with the name and contact info on top.
"""
    print("⏳ Generating CV using LLaMA 3.2 via Ollama API...")
    cv_text = generate_cv_with_llama(prompt)
    print("✅ CV content generated.\n")
    write_cv_to_word(cv_text)

if __name__ == "__main__":
    main()
